import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_resume_docx():
    doc = Document()
    
    # 标题
    title = doc.add_paragraph()
    run = title.add_run('个人简历 (兼职招聘/远程办公版)')
    run.bold = True
    run.font.size = Pt(18)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 基本信息
    doc.add_heading('基本信息', level=1)
    p = doc.add_paragraph()
    p.add_run('姓名/ID：').bold = True
    p.add_run('[请在此填写您的姓名/ID]\n')
    p.add_run('职位定位：').bold = True
    p.add_run('独立软件架构师 · 「天龙引擎」AI 交付引擎主理人\n')
    p.add_run('擅长领域：').bold = True
    p.add_run('企业级 Web 应用、Python 自动化生产力工具、AI 原生应用开发\n')
    p.add_run('服务承诺：').bold = True
    p.add_run('100% 交付确定性、80% 测试覆盖率、零死角架构文档')
    
    # 核心竞争优势
    doc.add_heading('💡 核心竞争优势：天龙引擎交付体系', level=1)
    doc.add_paragraph('我不仅仅是一名开发者，更是一套工业级交付标准的运营者。我主导研发的「天龙引擎」AI 辅助开发引擎，将软件工程从“手工作坊”升级为“自动化流水线”：')
    
    points = [
        ("防御性架构 (/01 Architect)", "项目启动即定义严格的 DoD (完成标准)，确保代码逻辑闭环，拒绝需求无限蔓延。"),
        ("极速精准构建 (/02 Builder)", "基于 React 18 & FastAPI 现代栈，通过自动化代码生成技术，比传统手动开发效率提升 3-5 倍。"),
        ("零 Bug 容错协议 (/03 Validator)", "强制执行自动化测试，覆盖所有核心业务路径，确保交付物处于“开箱即用”状态。"),
        ("安全与性能审计 (/04/05 Security & Reviewer)", "每一行代码必经深度审计，针对 API 泄露、SQL 注入、嵌套地狱进行专项排雷。"),
        ("文明级文档支撑 (/06 Scribe)", "自动化生成系统架构图、接口文档与维护说明，拒绝“代码跑路，文档失踪”。")
    ]
    
    for bold_text, normal_text in points:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(bold_text).bold = True
        p.add_run(f"：{normal_text}")
        
    # 技术底座
    doc.add_heading('🛠️ 技术底座 (Tech Stack)', level=1)
    tech_stack = [
        "前端：React 18 (TS 严格模式) | Tailwind CSS | 后台管理系统极速搭建",
        "后端：Python (FastAPI / Flask) | 数据结构化处理 | API 高并发设计",
        "工程化：Docker 容器化部署 | Git 规范化协同 | 自动化 CI/CD 测试流程",
        "AI 应用：大模型 API 集成 (LLM) | RAG 知识库构建 | 自动化 Agent 工作流"
    ]
    for tech in tech_stack:
        doc.add_paragraph(tech, style='List Bullet')
        
    # 代表性项目经验
    doc.add_heading('🌟 代表性项目经验 (基于天龙引擎交付)', level=1)
    
    p1 = doc.add_paragraph()
    p1.add_run('1. 企业级自动化数据处理与文档生成中心').bold = True
    doc.add_paragraph('通过代码自动处理复杂 Excel、Word 数据并转化为标准化可视化报告。实现了 100% 的操作自动化，将原本 2 小时的手动操作缩短至 1 分钟。', style='List Bullet')
    
    p2 = doc.add_paragraph()
    p2.add_run('2. 符合「防御性工程协议」的现代 Web Prototyping').bold = True
    doc.add_paragraph('引入天龙引擎审计流程，确保前端交互流畅（Loading/Error 状态全覆盖），后端 API 具备 30s 超时保护与错误自我修复逻辑。', style='List Bullet')
    
    # 交付承诺
    doc.add_heading('🛡️ 交付承诺 (Definition of Done)', level=1)
    dods = [
        "代码库：Clean Code 风格，遵循 Early Return 与防御性编程规范。",
        "测试集：全量通过的自动化测试脚本。",
        "文档包：包含系统部署手册、接口定义、数据字典。",
        "演示视频：完整的核心业务逻辑运行演示。"
    ]
    for dod in dods:
        doc.add_paragraph(dod, style='List Bullet')
        
    # 保存文档
    output_path = os.path.abspath('天龙引擎主理人_个人简历.docx')
    doc.save(output_path)
    print(f"Resume saved to: {output_path}")

if __name__ == "__main__":
    create_resume_docx()
