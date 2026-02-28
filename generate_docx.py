import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_platform_list():
    doc = Document()
    
    # 标题
    title = doc.add_heading('国内程序员兼职平台布局图 (2024-2025)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 定义梯队数据
    tiers = [
        {
            "name": "第一梯队：中高端、高质量（战略首选）",
            "platforms": [
                ["程序员客栈 (Proginn)", "国内领先的自由工作平台，中高端项目多", "https://www.proginn.com"],
                ["电鸭社区 (Eleduck)", "远程工作/外包项目社交社区，氛围极佳", "https://eleduck.com"],
                ["程聚宝 (原程序聚合)", "严选项目，人工审核，不收会员费", "https://www.chengjubao.com"]
            ]
        },
        {
            "name": "第二梯队：传统众包、快速获客（主力战场）",
            "platforms": [
                ["码市 (Codemart)", "CODING 旗下，需求透明，适合标准化开发", "https://codemart.com"],
                ["猿急送", "专为程序员定制，按需雇佣，交付效率高", "https://www.yuanjisong.com"],
                ["实现网", "专注于互联网人才兼职，时薪计费模式成熟", "https://shixian.com"],
                ["开源众包", "开源中国旗下，中小型软件定制需求多", "https://zb.oschina.net"]
            ]
        },
        {
            "name": "第三梯队：海量流量、长尾获客（尝试型渠道）",
            "platforms": [
                ["BOSS 直聘", "在意向中勾选“兼职/远程”，直接对话 BOSS", "https://www.zhipin.com"],
                ["猪八戒网 (ZBJ)", "综合型服务市场，适合新手和简单需求", "https://www.zbj.com"],
                ["闲鱼 / 小红书", "自定流量，发布“专业技术代工”服务", "APP 端搜索"]
            ]
        }
    ]
    
    for tier in tiers:
        # 梯队标题
        p = doc.add_paragraph()
        run = p.add_run(tier["name"])
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 51, 153) # 深蓝色
        
        # 创建表格
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        
        # 设置表头
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '平台名称'
        hdr_cells[1].text = '核心定位'
        hdr_cells[2].text = '官方网址'
        
        # 填充数据
        for plat in tier["platforms"]:
            row_cells = table.add_row().cells
            row_cells[0].text = plat[0]
            row_cells[1].text = plat[1]
            row_cells[2].text = plat[2]
            
        doc.add_paragraph() # 空行
        
    # 保存文档
    output_path = os.path.abspath('国内程序员兼职平台布局图.docx')
    doc.save(output_path)
    print(f"File saved to: {output_path}")

if __name__ == "__main__":
    create_platform_list()
